import os
import sys
import yaml
import markdown
import base64
import subprocess
import time
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_margins(cell, top, bottom, left, right):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, value in zip(['top', 'bottom', 'left', 'right'], [top, bottom, left, right]):
        if value is not None:
            node = OxmlElement(f'w:{margin}')
            node.set(qn('w:w'), str(value))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None, size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name, color in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if color:
            b = OxmlElement(f'w:{border_name}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), size)
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), color)
            tcBorders.append(b)
    tcPr.append(tcBorders)

def find_browser():
    candidates = [
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    ]
    for path in candidates:
        if os.path.exists(path):
            return path
    if sys.platform != "win32":
        try:
            return subprocess.check_output(["which", "google-chrome"]).decode().strip()
        except:
            pass
    return None

def parse_markdown(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    if content.startswith("---"):
        parts = content.split("---", 2)
        if len(parts) >= 3:
            front_matter = parts[1]
            markdown_body = parts[2]
            yaml_data = yaml.safe_load(front_matter) or {}
            return yaml_data, markdown_body
    
    return {}, content

def render_html_template(yaml_data, markdown_body, css_content, logo_base64, company_info):
    ref = yaml_data.get("reference", "")
    date = yaml_data.get("date", "")
    title = yaml_data.get("title", "Contrat")
    subtitle = yaml_data.get("subtitle", "")
    location = yaml_data.get("location", "Gamba")
    
    client = yaml_data.get("client", {})
    client_name = client.get("name", "")
    client_company = client.get("company", "")
    client_role = client.get("role", "")
    client_address = client.get("address", "")
    client_nif = client.get("nif", "")
    client_rccm = client.get("rccm", "")
    client_phone = client.get("phone", "")
    client_email = client.get("email", "")
    
    consultant = yaml_data.get("consultant", {})
    cons_name = consultant.get("name", "")
    cons_company = consultant.get("company", "")
    cons_role = consultant.get("role", "")
    cons_address = consultant.get("address", "")
    cons_phone = consultant.get("phone", "")
    cons_email = consultant.get("email", "")
    
    logo_src = f"data:image/png;base64,{logo_base64}" if logo_base64 else ""
    
    html_body = markdown.markdown(markdown_body, extensions=['extra', 'tables'])
    
    comp_name = company_info.get('name', 'IbogaLab')
    comp_html = f'<div style="font-size: 14pt; font-weight: bold; color: var(--primary-color); margin-bottom: 5px;">{comp_name}</div>'
    if company_info.get('address'): comp_html += f"<div>{company_info['address'].strip()}</div>"
    if company_info.get('email'): comp_html += f"<div>{company_info['email'].strip()}</div>"
    if company_info.get('phone'): comp_html += f"<div>Tél: {company_info['phone'].strip()}</div>"
    if company_info.get('website'): comp_html += f"<div>{company_info['website'].strip()}</div>"
    
    html = f"""<!DOCTYPE html>
<html lang="fr">
<head>
  <meta charset="UTF-8">
  <title>{title} - {ref}</title>
  <style>
    {css_content}
  </style>
</head>
<body>
  <div class="invoice-container">
    <div class="header" style="display: flex; justify-content: space-between; align-items: flex-start;">
      <div>
        <img class="company-logo" src="{logo_src}" alt="Logo">
      </div>
      <div style="text-align: right; font-size: 10pt; color: #718096; line-height: 1.4;">
        {comp_html}
      </div>
    </div>
    
    <div class="document-title-bar">
      <div class="title-left">{title}</div>
      <div class="meta-right">Réf: {ref}</div>
    </div>
    
    <div style="text-align: center; margin-bottom: 20px;">
        <h2 style="font-size: 14pt; margin-bottom: 5px; color: var(--primary-color);">{subtitle}</h2>
        <div style="font-size: 10pt; color: var(--text-light);">Fait à {location}, le {date}</div>
    </div>
    
    <div style="font-size: 10pt; font-weight: bold; text-align: center; margin: 30px 0 15px 0; text-transform: uppercase;">
        Entre les soussignés
    </div>
    
    <div class="details-grid">
      <div class="details-card">
        <h3>Le Client</h3>
        <div class="details-line"><strong>{client_company}</strong></div>
        <div class="details-line">Représenté par : {client_name}</div>
        <div class="details-line">Qualité : {client_role}</div>
        <div class="details-line">Adresse : {client_address}</div>
"""
    if client_nif: html += f'        <div class="details-line">NIF : {client_nif}</div>\n'
    if client_rccm: html += f'        <div class="details-line">RCCM : {client_rccm}</div>\n'
    if client_phone: html += f'        <div class="details-line">Tél : {client_phone}</div>\n'
    if client_email: html += f'        <div class="details-line">Email : {client_email}</div>\n'
    
    html += f"""      </div>
      <div class="details-card">
        <h3>Le Consultant</h3>
        <div class="details-line"><strong>{cons_name}</strong></div>
        <div class="details-line">Structure : {cons_company}</div>
        <div class="details-line">Qualité : {cons_role}</div>
        <div class="details-line">Adresse : {cons_address}</div>
"""
    if cons_phone: html += f'        <div class="details-line">Tél : {cons_phone}</div>\n'
    if cons_email: html += f'        <div class="details-line">Email : {cons_email}</div>\n'
    
    html += f"""      </div>
    </div>
    
    <div class="contract-body">
      {html_body}
    </div>
    
    <div class="signatures-section">
      <div class="signature-box">
        <div class="title">Pour le Client</div>
        <div class="label-sub">Lu et approuvé</div>
        <div style="margin-top: 120px;">
            <strong>{client_name}</strong><br>
            <span style="font-size: 8pt; color: #718096;">{client_role}</span>
        </div>
      </div>
      <div class="signature-box">
        <div class="title">Pour le Consultant</div>
        <div class="label-sub">Lu et approuvé</div>
        <div style="margin-top: 120px;">
            <strong>{cons_name}</strong><br>
            <span style="font-size: 8pt; color: #718096;">{cons_role}</span>
        </div>
      </div>
    </div>
    
  </div>
</body>
</html>
"""
    return html

def generate_docx(yaml_data, markdown_body, logo_path, docx_path, company_info):
    ref = yaml_data.get("reference", "")
    date = yaml_data.get("date", "")
    title = yaml_data.get("title", "Contrat")
    subtitle = yaml_data.get("subtitle", "")
    location = yaml_data.get("location", "Gamba")
    
    client = yaml_data.get("client", {})
    consultant = yaml_data.get("consultant", {})
    
    doc = Document()
    
    c_forest = RGBColor(16, 56, 36)      # #103824
    c_sage = RGBColor(102, 157, 105)    # #669D69
    c_dark = RGBColor(45, 55, 72)       # #2D3748
    c_light = RGBColor(113, 128, 150)   # #718096
    
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = c_dark
    
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    header_table.columns[0].width = Inches(3.8)
    header_table.columns[1].width = Inches(3.2)
    
    left_cell = header_table.cell(0, 0)
    right_cell = header_table.cell(0, 1)
    
    p_logo = left_cell.paragraphs[0]
    if os.path.exists(logo_path):
        p_logo.add_run().add_picture(logo_path, width=Inches(1.4))
        
    p_comp = right_cell.paragraphs[0]
    p_comp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_name = p_comp.add_run(f"{company_info.get('name', 'IbogaLab')}\\n")
    run_name.bold = True
    run_name.font.color.rgb = c_forest
    run_name.font.size = Pt(11.5)
    
    comp_lines = []
    if company_info.get('address'): comp_lines.append(company_info['address'].strip())
    if company_info.get('email'): comp_lines.append(company_info['email'].strip())
    if company_info.get('phone'): comp_lines.append(f"Tél: {company_info['phone'].strip()}")
    if company_info.get('website'): comp_lines.append(company_info['website'].strip())
    
    comp_details_text = "\\n".join(comp_lines)
    run_info = p_comp.add_run(comp_details_text)
    run_info.font.color.rgb = c_light
    run_info.font.size = Pt(8)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    title_bar_table = doc.add_table(rows=1, cols=2)
    title_bar_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    title_bar_table.autofit = False
    title_bar_table.columns[0].width = Inches(3.5)
    title_bar_table.columns[1].width = Inches(3.2)
    
    c_title = title_bar_table.cell(0, 0)
    c_meta = title_bar_table.cell(0, 1)
    set_cell_background(c_title, "103824")
    set_cell_background(c_meta, "103824")
    set_cell_margins(c_title, 100, 100, 150, 150)
    set_cell_margins(c_meta, 100, 100, 150, 150)
    
    p_title = c_title.paragraphs[0]
    run_title_text = p_title.add_run(title.upper())
    run_title_text.bold = True
    run_title_text.font.size = Pt(10)
    run_title_text.font.color.rgb = RGBColor(255, 255, 255)
    
    p_meta = c_meta.paragraphs[0]
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_meta_text = p_meta.add_run(f"Réf: {ref}")
    run_meta_text.font.size = Pt(8.5)
    run_meta_text.font.color.rgb = RGBColor(255, 255, 255)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(15)
    p_sub.paragraph_format.space_after = Pt(2)
    run_sub = p_sub.add_run(subtitle)
    run_sub.bold = True
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = c_forest
    
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run(f"Fait à {location}, le {date}")
    run_date.font.size = Pt(9.5)
    run_date.font.color.rgb = c_light
    
    p_entre = doc.add_paragraph()
    p_entre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_entre.paragraph_format.space_before = Pt(20)
    p_entre.paragraph_format.space_after = Pt(10)
    run_entre = p_entre.add_run("ENTRE LES SOUSSIGNÉS")
    run_entre.bold = True
    
    info_table = doc.add_table(rows=1, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    info_table.columns[0].width = Inches(3.3)
    info_table.columns[1].width = Inches(3.4)
    
    cell_client = info_table.cell(0, 0)
    cell_cons = info_table.cell(0, 1)
    
    set_cell_background(cell_client, "FFFFFF")
    set_cell_background(cell_cons, "FFFFFF")
    set_cell_margins(cell_client, 120, 120, 150, 150)
    set_cell_margins(cell_cons, 120, 120, 150, 150)
    set_cell_borders(cell_client, "669D69", "669D69", "669D69", "669D69", "4")
    set_cell_borders(cell_cons, "669D69", "669D69", "669D69", "669D69", "4")
    
    p_c = cell_client.paragraphs[0]
    run_c_hdr = p_c.add_run("LE CLIENT\n")
    run_c_hdr.bold = True
    run_c_hdr.font.size = Pt(8.5)
    run_c_hdr.font.color.rgb = c_sage
    
    c_lines = [f"**{client.get('company', '')}**", f"Représenté par : {client.get('name', '')}", f"Qualité : {client.get('role', '')}", f"Adresse : {client.get('address', '')}"]
    if client.get('nif'): c_lines.append(f"NIF : {client['nif']}")
    if client.get('rccm'): c_lines.append(f"RCCM : {client['rccm']}")
    if client.get('phone'): c_lines.append(f"Tél : {client['phone']}")
    if client.get('email'): c_lines.append(f"Email : {client['email']}")
    
    for line in c_lines:
        if line.startswith("**"):
            r = p_c.add_run(line.replace("**", "") + "\n")
            r.bold = True
        else:
            r = p_c.add_run(line + "\n")
        r.font.size = Pt(9)
        
    p_cons = cell_cons.paragraphs[0]
    run_cons_hdr = p_cons.add_run("LE CONSULTANT\n")
    run_cons_hdr.bold = True
    run_cons_hdr.font.size = Pt(8.5)
    run_cons_hdr.font.color.rgb = c_sage
    
    cons_lines = [f"**{consultant.get('name', '')}**", f"Structure : {consultant.get('company', '')}", f"Qualité : {consultant.get('role', '')}", f"Adresse : {consultant.get('address', '')}"]
    if consultant.get('phone'): cons_lines.append(f"Tél : {consultant['phone']}")
    if consultant.get('email'): cons_lines.append(f"Email : {consultant['email']}")
    
    for line in cons_lines:
        if line.startswith("**"):
            r = p_cons.add_run(line.replace("**", "") + "\n")
            r.bold = True
        else:
            r = p_cons.add_run(line + "\n")
        r.font.size = Pt(9)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(15)
    
    # Process markdown body
    for line in markdown_body.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(15)
            p.paragraph_format.space_after = Pt(5)
            r = p.add_run(line[2:])
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = c_forest
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(line[3:])
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = c_sage
        elif line.startswith('• ') or line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            # Remove the markdown list marker and replace with a nice bullet point
            clean_line = line[2:].strip()
            r = p.add_run('•  ' + clean_line)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Basic bold rendering for markdown
            parts = line.split('**')
            for i, part in enumerate(parts):
                if part:
                    r = p.add_run(part)
                    if i % 2 != 0:
                        r.bold = True
                        
    doc.add_paragraph().paragraph_format.space_after = Pt(30)
    
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False
    sig_table.columns[0].width = Inches(3.3)
    sig_table.columns[1].width = Inches(3.4)
    
    c_sig_client = sig_table.cell(0, 0)
    c_sig_cons = sig_table.cell(0, 1)
    
    set_cell_borders(c_sig_client, top="E2E8F0", bottom=None, left=None, right=None, size="4")
    set_cell_borders(c_sig_cons, top="E2E8F0", bottom=None, left=None, right=None, size="4")
    set_cell_margins(c_sig_client, 100, 100, 100, 100)
    set_cell_margins(c_sig_cons, 100, 100, 100, 100)
    
    p_sig_c = c_sig_client.paragraphs[0]
    p_sig_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sig_c.paragraph_format.space_before = Pt(10)
    r1 = p_sig_c.add_run("Pour le Client\nLu et approuvé\n\n\n\n\n\n\n\n")
    r1.font.size = Pt(9)
    r2 = p_sig_c.add_run(client.get('name', ''))
    r2.bold = True
    r2.font.size = Pt(9)
    r3 = p_sig_c.add_run("\n" + client.get('role', ''))
    r3.font.size = Pt(8)
    r3.font.color.rgb = c_light
    
    p_sig_a = c_sig_cons.paragraphs[0]
    p_sig_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sig_a.paragraph_format.space_before = Pt(10)
    r4 = p_sig_a.add_run("Pour le Consultant\nLu et approuvé\n\n\n\n\n\n\n\n")
    r4.font.size = Pt(9)
    r5 = p_sig_a.add_run(consultant.get('name', ''))
    r5.bold = True
    r5.font.size = Pt(9)
    r6 = p_sig_a.add_run("\n" + consultant.get('role', ''))
    r6.font.size = Pt(8)
    r6.font.color.rgb = c_light
    
    doc.save(docx_path)

def compile_document(md_path):
    print(f"Compiling {md_path}...")
    yaml_data, markdown_body = parse_markdown(md_path)
    
    dir_name = os.path.dirname(os.path.abspath(md_path))
    base_name = os.path.splitext(os.path.basename(md_path))[0]
    
    html_path = os.path.join(dir_name, f"{base_name}.html")
    pdf_path = os.path.join(dir_name, f"{base_name}.pdf")
    docx_path = os.path.join(dir_name, f"{base_name}.docx")
    
    skill_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    css_file = os.path.join(skill_dir, "resources", "style.css")
    css_content = ""
    if os.path.exists(css_file):
        with open(css_file, 'r', encoding='utf-8') as f:
            css_content = f.read()
            
    logo_path = os.path.join(skill_dir, "resources", "logo1.png")
    logo_base64 = ""
    if os.path.exists(logo_path):
        with open(logo_path, 'rb') as f:
            logo_base64 = base64.b64encode(f.read()).decode('utf-8')
            
    company_info = {
        "name": "IbogaLab",
        "address": "Libreville, Gabon",
        "email": "contact@ibogalab.tech",
        "phone": "+241 066195786",
        "website": "www.ibogalab.tech"
    }
    
    settings_path = os.path.expanduser("~/.ibogalab-invoice-settings.json")
    if os.path.exists(settings_path):
        import json
        with open(settings_path, 'r', encoding='utf-8') as f:
            try:
                user_settings = json.load(f)
                if "company" in user_settings:
                    company_info.update(user_settings["company"])
            except:
                pass
            
    html_content = render_html_template(yaml_data, markdown_body, css_content, logo_base64, company_info)
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    print(f"Generated HTML: {html_path}")
    
    browser_path = find_browser()
    if browser_path:
        print(f"Using browser: {browser_path} to print to PDF...")
        try:
            html_url = "file:///" + os.path.abspath(html_path).replace("\\", "/")
            subprocess.run([
                browser_path, 
                "--headless", 
                "--disable-gpu", 
                "--no-pdf-header-footer",
                f"--print-to-pdf={pdf_path}", 
                html_url
            ], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            print(f"Generated PDF: {pdf_path}")
        except Exception as e:
            print(f"Failed to generate PDF: {e}")
            
    generate_docx(yaml_data, markdown_body, logo_path, docx_path, company_info)
    print(f"Generated DOCX: {docx_path}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python sync_contract.py <path_to_markdown>")
        sys.exit(1)
    file_path = sys.argv[-1]
    compile_document(file_path)
