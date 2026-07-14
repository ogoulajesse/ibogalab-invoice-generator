import os
import sys
import re
import json
import base64
import time
import shutil
import subprocess

# Self-install pyyaml if missing
try:
    import yaml
except ImportError:
    print("pyyaml is missing, installing...")
    subprocess.run([sys.executable, "-m", "pip", "install", "pyyaml", "--quiet"])
    import yaml

try:
    import markdown
except ImportError:
    print("markdown is missing, installing...")
    subprocess.run([sys.executable, "-m", "pip", "install", "markdown", "--quiet"])
    import markdown

try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("python-docx is missing, installing...")
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "--quiet"])
    import docx
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn

SETTINGS_FILE = os.path.expanduser("~/.ibogalab-invoice-settings.json")

DEFAULT_SETTINGS = {
    "counters": {
        "IBGL-2026-D": 1,
        "IBGL-2026-F": 1
    },
    "company": {
        "name": "IbogaLab",
        "address": "Libreville, Gabon",
        "email": "contact@ibogalab.com",
        "rccm": "[Votre Numéro]",
        "nif": "[Votre NIF]",
        "capital": "1 000 000 FCFA",
        "website": "www.ibogalab.tech"
    }
}

def load_settings():
    settings = json.loads(json.dumps(DEFAULT_SETTINGS))
    if os.path.exists(SETTINGS_FILE):
        try:
            with open(SETTINGS_FILE, 'r', encoding='utf-8') as f:
                user_settings = json.load(f)
                if "counters" in user_settings:
                    settings["counters"].update(user_settings["counters"])
                if "company" in user_settings:
                    settings["company"].update(user_settings["company"])
        except Exception:
            pass
    else:
        save_settings(settings)
    return settings

def save_settings(settings):
    try:
        with open(SETTINGS_FILE, 'w', encoding='utf-8') as f:
            json.dump(settings, f, indent=2, ensure_ascii=False)
    except Exception as e:
        print(f"Warning: could not save settings to {SETTINGS_FILE}: {e}")

def get_next_number(doc_type, prefix=None):
    settings = load_settings()
    year = time.strftime("%Y")
    
    if not prefix:
        if doc_type.upper() == "DEVIS":
            prefix = f"IBGL-{year}-D"
        else:
            prefix = f"IBGL-{year}-F"
            
    counters = settings.setdefault("counters", {})
    index = counters.setdefault(prefix, 1)
    
    doc_number = f"{prefix}{index:02d}"
    
    counters[prefix] = index + 1
    save_settings(settings)
    
    return doc_number

def format_currency(val, currency="FCFA"):
    try:
        val_int = int(round(float(val)))
        return f"{val_int:,}".replace(",", " ") + f" {currency}"
    except Exception:
        return f"{val} {currency}"

def find_browser():
    paths = [
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    ]
    for path in paths:
        if os.path.exists(path):
            return path
    for cmd in ["chrome", "msedge"]:
        path = shutil.which(cmd)
        if path:
            return path
    return None

# Helper functions for python-docx styling
def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, top="CBD3CE", bottom="CBD3CE", left="CBD3CE", right="CBD3CE", size="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, color in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if color:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), size)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
            tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_bottom_border(cell, hex_color="CBD3CE", size="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), hex_color)
    tcBorders.append(bottom)
    tcPr.append(tcBorders)

def parse_markdown(md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
        
    yaml_data = {}
    markdown_body = content
    
    if content.startswith("---"):
        parts = content.split("---", 2)
        if len(parts) >= 3:
            try:
                yaml_data = yaml.safe_load(parts[1])
                markdown_body = parts[2]
            except Exception as e:
                print(f"Error parsing YAML front-matter: {e}")
                
    return yaml_data, markdown_body, content

def update_markdown_number(md_path, old_content, new_number):
    pattern1 = r'(number:\s*[\'"]?)(auto)([\'"]?)'
    if re.search(pattern1, old_content):
        updated_content = re.sub(pattern1, rf'\g<1>{new_number}\g<3>', old_content, count=1)
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(updated_content)
        return updated_content
    return old_content

def process_items_table(markdown_body, auto_calculate=True, currency="FCFA", tax_rate=0.0):
    lines = markdown_body.strip().split('\n')
    table_lines = []
    in_table = False
    
    for line in lines:
        if line.strip().startswith('|'):
            in_table = True
            table_lines.append(line.strip())
        elif in_table:
            break
            
    if not table_lines:
        return [], 0, 0, 0, ""
        
    headers = [h.strip() for h in table_lines[0].split('|')[1:-1]]
    rows = []
    for line in table_lines[2:]:
        if line.strip():
            rows.append([cell.strip() for cell in line.split('|')[1:-1]])
            
    idx_desc = 0
    idx_unit = 1
    idx_qty = 2
    idx_pu = 3
    
    table_data = []
    total_ht = 0
    
    for row in rows:
        if len(row) < 4:
            continue
        desc = row[idx_desc]
        unit = row[idx_unit]
        
        try:
            qty_str = re.sub(r'[^\d.]', '', row[idx_qty])
            qty = float(qty_str) if '.' in qty_str else int(qty_str)
        except Exception:
            qty = 1
            
        try:
            pu_str = re.sub(r'[^\d.]', '', row[idx_pu])
            pu = float(pu_str) if '.' in pu_str else int(pu_str)
        except Exception:
            pu = 0
            
        amount = qty * pu
        total_ht += amount
        
        table_data.append({
            "designation": desc,
            "unite": unit,
            "qty": qty,
            "pu": pu,
            "amount": amount
        })
        
    new_table_md = "| Désignation & Détails | Unité | Qté | P.U. HT | Montant HT |\n"
    new_table_md += "| :--- | :---: | :---: | :---: | :---: |\n"
    for item in table_data:
        pu_f = format_currency(item["pu"], currency)
        amount_f = format_currency(item["amount"], currency)
        desc_md = item["designation"]
        new_table_md += f"| {desc_md} | {item['unite']} | {item['qty']} | {pu_f} | {amount_f} |\n"
        
    vat_amount = total_ht * tax_rate
    total_ttc = total_ht + vat_amount
    
    return table_data, total_ht, vat_amount, total_ttc, new_table_md

def compile_document(md_path):
    print(f"Compiling {md_path}...")
    yaml_data, markdown_body, raw_content = parse_markdown(md_path)
    
    doc_type = yaml_data.get("type", "DEVIS").upper()
    doc_number = yaml_data.get("number", "auto")
    
    if doc_number == "auto":
        prefix = yaml_data.get("prefix", None)
        doc_number = get_next_number(doc_type, prefix)
        raw_content = update_markdown_number(md_path, raw_content, doc_number)
        yaml_data["number"] = doc_number
        print(f"Assigned new number: {doc_number}")
        
    currency = yaml_data.get("currency", "FCFA")
    tax_rate = float(yaml_data.get("tax_rate", 0.0))
    auto_calc = yaml_data.get("auto_calculate", True)
    
    items_data, total_ht, vat_amount, total_ttc, table_md = process_items_table(
        markdown_body, auto_calc, currency, tax_rate
    )
    
    settings = load_settings()
    company_info = settings.get("company", DEFAULT_SETTINGS["company"])
    
    dir_name = os.path.dirname(os.path.abspath(md_path))
    base_name = os.path.splitext(os.path.basename(md_path))[0]
    
    html_path = os.path.join(dir_name, f"{base_name}.html")
    pdf_path = os.path.join(dir_name, f"{base_name}.pdf")
    docx_path = os.path.join(dir_name, f"{base_name}.docx")
    
    skill_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    css_file = os.path.join(skill_dir, "resources", "style.css")
    css_content = ""
    if os.path.exists(css_file):
        with open(css_file, 'r', encoding='utf-8') as f:
            css_content = f.read()
            
    logo_path = os.path.join(skill_dir, "resources", "logo1.png")
    logo_base64 = ""
    if os.path.exists(logo_path):
        with open(logo_path, 'rb') as f:
            logo_base64 = base64.b64encode(f.read()).decode('utf-8')
            
    html_table = markdown.markdown(table_md, extensions=['extra', 'tables'])
    
    html_content = render_html_template(yaml_data, html_table, total_ht, vat_amount, total_ttc, css_content, logo_base64, company_info)
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    print(f"Generated HTML: {html_path}")
    
    browser_path = find_browser()
    if browser_path:
        print(f"Using browser: {browser_path} to print to PDF (no header/footer)...")
        try:
            html_url = "file:///" + os.path.abspath(html_path).replace("\\", "/")
            subprocess.run([
                browser_path, 
                "--headless", 
                "--disable-gpu", 
                "--no-pdf-header-footer",
                f"--print-to-pdf={pdf_path}", 
                html_url
            ], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            print(f"Generated PDF: {pdf_path}")
        except Exception as e:
            print(f"Failed to generate PDF: {e}")
    else:
        print("Error: Could not find Google Chrome or Microsoft Edge to generate PDF.")
        
    generate_docx(yaml_data, items_data, total_ht, vat_amount, total_ttc, logo_path, docx_path, company_info)
    print(f"Generated DOCX: {docx_path}")
    
    return html_path, pdf_path, docx_path

def render_html_template(yaml_data, html_table, total_ht, vat_amount, total_ttc, css_content, logo_base64, company_info):
    doc_type = yaml_data.get("type", "DEVIS").upper()
    doc_number = yaml_data.get("number", "N/A")
    doc_date = yaml_data.get("date", "N/A")
    doc_validity = yaml_data.get("validity", "")
    due_date = yaml_data.get("due_date", "")
    currency = yaml_data.get("currency", "FCFA")
    
    client = yaml_data.get("client", {})
    client_name = client.get("name", "")
    client_company = client.get("company", "")
    client_address = client.get("address", "")
    
    project = yaml_data.get("project", {})
    proj_name = project.get("name", "")
    proj_desc = project.get("description", "")
    proj_duration = project.get("duration", "")
    ref_devis = yaml_data.get("reference_devis", "")
    
    logo_src = f"data:image/png;base64,{logo_base64}" if logo_base64 else ""
    
    meta_date_text = f"Établi le {doc_date}"
    if doc_type == "DEVIS" and doc_validity:
        meta_date_text += f" - Valide {doc_validity}"
    elif doc_type == "FACTURE" and due_date:
        meta_date_text += f" - Échéance {due_date}"
        
    client_lines = []
    if client_name: client_lines.append(f"<div>{client_name}</div>")
    if client_company: client_lines.append(f"<div>{client_company}</div>")
    if client_address: client_lines.append(f"<div>{client_address}</div>")
    client_html = "\n".join(client_lines)
    
    project_lines = []
    if proj_name: project_lines.append(f"<div>{proj_name}</div>")
    if proj_desc: project_lines.append(f"<div>{proj_desc}</div>")
    if proj_duration: project_lines.append(f"<div>Délai estimé : {proj_duration}</div>")
    if doc_type == "FACTURE" and ref_devis:
         project_lines.append(f"<div>Réf Devis : {ref_devis}</div>")
    project_html = "\n".join(project_lines)
    
    recurring_html = ""
    recurring_costs = yaml_data.get("recurring_costs", [])
    if recurring_costs:
        recurring_html += "<div class='extra-section'>"
        for item in recurring_costs:
            recurring_html += f"""
            <div class='extra-item'>
              <span>Estimation mensuelle post-lancement : {item.get('designation', '')}</span>
              <strong>{item.get('price', '')} / {item.get('period', 'mois')}</strong>
            </div>
            """
        recurring_html += "</div>"
        
    total_ht_f = format_currency(total_ht, currency)
    total_ttc_val = f"{int(round(total_ttc)):,}".replace(",", " ")
    
    doc_type_label = "Développement" if doc_type == "DEVIS" else "Facture"
    summary_html = f"""
    <div class="summary-block">
      <div class="summary-row-ht">
        <span>Total {doc_type_label} HT</span>
        <span>{total_ht_f}</span>
      </div>
    """
    if vat_amount > 0:
        summary_html += f"""
      <div class="summary-row-ht">
        <span>TVA ({int(tax_rate*100)}%)</span>
        <span>{format_currency(vat_amount, currency)}</span>
      </div>
        """
    summary_html += f"""
      <div class="totals-divider"></div>
      <div class="summary-row-ttc">
        <span>Total Général TTC</span>
        <div class="amount-wrapper">
          {total_ttc_val}
          <span class="amount-currency">{currency}</span>
        </div>
      </div>
    </div>
    """
    
    terms_html = ""
    terms = yaml_data.get("terms", [])
    if terms:
        terms_html += "<div class='terms-section'>"
        terms_html += f"<h3>Conditions de paiement &amp; Modalités :</h3>"
        terms_html += "<ul>"
        for term in terms:
            terms_html += f"<li>{term}</li>"
        terms_html += "</ul>"
        terms_html += "</div>"
        
    signatures_html = f"""
    <div class="signatures-section">
      <div class="signature-box">
        <div class="title">Bon pour accord</div>
        <div class="label-sub">Signature du Client (précédée de la date)</div>
      </div>
      <div class="signature-box">
        <div class="title">Cachet et Signature de l'Agence</div>
      </div>
    </div>
    """
    
    return f"""<!DOCTYPE html>
<html lang="fr">
<head>
  <meta charset="UTF-8">
  <title>{doc_type} N° {doc_number} - {client_company}</title>
  <style>
    {css_content}
  </style>
</head>
<body>
  <div class="invoice-container">
    <div class="header">
      <div>
        <img class="company-logo" src="{logo_src}" alt="Iboga Lab Logo">
      </div>
      <div class="company-info" style="text-align: right;">
        <div class="company-name">{company_info.get('name', 'IbogaLab')}</div>
        <div>{company_info.get('address', 'Libreville, Gabon')}</div>
        <div>{company_info.get('email', 'contact@ibogalab.com')}</div>
        <div>RCCM: {company_info.get('rccm', '[Votre Numéro]')}</div>
        <div>NIF: {company_info.get('nif', '[Votre NIF]')}</div>
      </div>
    </div>
    
    <div class="document-title-bar">
      <div class="title-left">{doc_type} N° {doc_number}</div>
      <div class="meta-right">{meta_date_text}</div>
    </div>
    
    <div class="details-grid">
      <div class="details-card">
        <h3>Client</h3>
        <div class="details-line">{client_html}</div>
      </div>
      <div class="details-card">
        <h3>Projet</h3>
        <div class="details-line">{project_html}</div>
      </div>
    </div>
    
    <div class="items-table-wrapper">
      {html_table}
    </div>
    
    {recurring_html}
    {summary_html}
    {terms_html}
    {signatures_html}
    
    <div class="footer-legal">
      {company_info.get('name', 'Iboga Lab')} SARL - Capital Social de {company_info.get('capital', '1 000 000 FCFA')} - {company_info.get('address', 'Libreville, Gabon')} - {company_info.get('website', 'www.ibogalab.tech')}
    </div>
  </div>
</body>
</html>
"""

def generate_docx(yaml_data, items_data, total_ht, vat_amount, total_ttc, logo_path, docx_path, company_info):
    doc_type = yaml_data.get("type", "DEVIS").upper()
    doc_number = yaml_data.get("number", "N/A")
    doc_date = yaml_data.get("date", "N/A")
    doc_validity = yaml_data.get("validity", "")
    due_date = yaml_data.get("due_date", "")
    currency = yaml_data.get("currency", "FCFA")
    
    client = yaml_data.get("client", {})
    client_name = client.get("name", "")
    client_company = client.get("company", "")
    client_address = client.get("address", "")
    
    project = yaml_data.get("project", {})
    proj_name = project.get("name", "")
    proj_desc = project.get("description", "")
    proj_duration = project.get("duration", "")
    ref_devis = yaml_data.get("reference_devis", "")
    
    doc = Document()
    
    c_forest = RGBColor(16, 56, 36)      # #103824
    c_sage = RGBColor(102, 157, 105)    # #669D69
    c_dark = RGBColor(45, 55, 72)       # #2D3748
    c_light = RGBColor(113, 128, 150)   # #718096
    
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(9.5)
    normal_style.font.color.rgb = c_dark
    
    # Header: Logo & Company details
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    header_table.columns[0].width = Inches(3.8)
    header_table.columns[1].width = Inches(3.2)
    
    left_cell = header_table.cell(0, 0)
    right_cell = header_table.cell(0, 1)
    
    p_logo = left_cell.paragraphs[0]
    if os.path.exists(logo_path):
        p_logo.add_run().add_picture(logo_path, width=Inches(1.4))
        
    p_comp = right_cell.paragraphs[0]
    p_comp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_name = p_comp.add_run(f"{company_info.get('name', 'IbogaLab')}\n")
    run_name.bold = True
    run_name.font.color.rgb = c_forest
    run_name.font.size = Pt(11.5)
    
    comp_details_text = f"{company_info.get('address', 'Libreville, Gabon')}\n{company_info.get('email', 'contact@ibogalab.com')}\nRCCM: {company_info.get('rccm', '[Votre Numéro]')}\nNIF: {company_info.get('nif', '[Votre NIF]')}"
    run_info = p_comp.add_run(comp_details_text)
    run_info.font.color.rgb = c_light
    run_info.font.size = Pt(8)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    title_bar_table = doc.add_table(rows=1, cols=2)
    title_bar_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    title_bar_table.autofit = False
    title_bar_table.columns[0].width = Inches(3.5)
    title_bar_table.columns[1].width = Inches(3.5)
    
    c_title = title_bar_table.cell(0, 0)
    c_meta = title_bar_table.cell(0, 1)
    set_cell_background(c_title, "103824")
    set_cell_background(c_meta, "103824")
    set_cell_margins(c_title, 100, 100, 150, 150)
    set_cell_margins(c_meta, 100, 100, 150, 150)
    
    p_title = c_title.paragraphs[0]
    run_title_text = p_title.add_run(f"{doc_type} N° {doc_number}")
    run_title_text.bold = True
    run_title_text.font.size = Pt(10.5)
    run_title_text.font.color.rgb = RGBColor(255, 255, 255)
    
    p_meta = c_meta.paragraphs[0]
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta_date_text = f"Établi le {doc_date}"
    if doc_type == "DEVIS" and doc_validity:
        meta_date_text += f" - Valide {doc_validity}"
    elif doc_type == "FACTURE" and due_date:
        meta_date_text += f" - Échéance {due_date}"
    run_meta_text = p_meta.add_run(meta_date_text)
    run_meta_text.font.size = Pt(8.5)
    run_meta_text.font.color.rgb = RGBColor(255, 255, 255)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(15)
    
    info_table = doc.add_table(rows=1, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    info_table.columns[0].width = Inches(3.4)
    info_table.columns[1].width = Inches(3.6)
    
    cell_client = info_table.cell(0, 0)
    cell_project = info_table.cell(0, 1)
    
    set_cell_background(cell_client, "FFFFFF")
    set_cell_background(cell_project, "FFFFFF")
    set_cell_margins(cell_client, 120, 120, 150, 150)
    set_cell_margins(cell_project, 120, 120, 150, 150)
    
    set_cell_borders(cell_client, "669D69", "669D69", "669D69", "669D69", "4")
    set_cell_borders(cell_project, "669D69", "669D69", "669D69", "669D69", "4")
    
    p_c = cell_client.paragraphs[0]
    run_c_hdr = p_c.add_run("CLIENT\n")
    run_c_hdr.bold = True
    run_c_hdr.font.size = Pt(8)
    run_c_hdr.font.color.rgb = c_sage
    
    run_c_det = p_c.add_run(f"{client_name}\n{client_company}\n{client_address}")
    run_c_det.font.size = Pt(9)
    
    p_p = cell_project.paragraphs[0]
    run_p_hdr = p_p.add_run("PROJET\n")
    run_p_hdr.bold = True
    run_p_hdr.font.size = Pt(8)
    run_p_hdr.font.color.rgb = c_sage
    
    proj_details = f"{proj_name}\n{proj_desc}"
    if proj_duration:
        proj_details += f"\nDélai estimé : {proj_duration}"
    if doc_type == "FACTURE" and ref_devis:
        proj_details += f"\nRéf Devis : {ref_devis}"
    run_p_det = p_p.add_run(proj_details)
    run_p_det.font.size = Pt(9)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(15)
    
    items_table = doc.add_table(rows=1, cols=5)
    items_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    items_table.autofit = False
    
    col_widths = [Inches(3.3), Inches(0.8), Inches(0.5), Inches(1.1), Inches(1.3)]
    for idx, width in enumerate(col_widths):
        items_table.columns[idx].width = width
        
    hdr_cells = items_table.rows[0].cells
    headers = ["Désignation & Détails", "Unité", "Qté", "P.U. HT", "Montant HT"]
    for idx, name in enumerate(headers):
        hdr_cells[idx].text = name
        set_cell_background(hdr_cells[idx], "103824")
        set_cell_margins(hdr_cells[idx], 120, 120, 120, 120)
        p = hdr_cells[idx].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(8.5)
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        if idx in [1, 2]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif idx in [3, 4]:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
    for r_idx, item in enumerate(items_data):
        row_cells = items_table.add_row().cells
        
        for idx in range(5):
            set_cell_background(row_cells[idx], "FFFFFF")
            set_cell_margins(row_cells[idx], 120, 120, 120, 120)
            set_cell_bottom_border(row_cells[idx], "E2E8F0", "4")
            
        desc_cell = row_cells[0]
        desc_p = desc_cell.paragraphs[0]
        desc_p.paragraph_format.space_before = Pt(0)
        desc_p.paragraph_format.space_after = Pt(2)
        
        full_desc = item["designation"]
        if "<br>" in full_desc:
            title, details = full_desc.split("<br>", 1)
        elif "\n" in full_desc:
            title, details = full_desc.split("\n", 1)
        else:
            title, details = full_desc, ""
            
        title = title.replace("**", "").replace("__", "").strip()
        details = details.replace("**", "").replace("__", "").strip()
        
        run_title = desc_p.add_run(title)
        run_title.bold = True
        run_title.font.size = Pt(9.5)
        run_title.font.color.rgb = c_forest
        
        if details:
            desc_p_sub = desc_cell.add_paragraph()
            desc_p_sub.paragraph_format.space_before = Pt(0)
            desc_p_sub.paragraph_format.space_after = Pt(0)
            run_desc = desc_p_sub.add_run(details)
            run_desc.font.size = Pt(8.5)
            run_desc.font.color.rgb = c_light
            
        row_cells[1].text = str(item["unite"])
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].text = str(item["qty"])
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        row_cells[3].text = format_currency(item["pu"], currency)
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        row_cells[4].text = format_currency(item["amount"], currency)
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        for idx in [1, 2, 3, 4]:
            p = row_cells[idx].paragraphs[0]
            if p.runs:
                p.runs[0].font.size = Pt(9)
                
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    recurring_costs = yaml_data.get("recurring_costs", [])
    if recurring_costs:
        rec_table = doc.add_table(rows=0, cols=2)
        rec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        rec_table.autofit = False
        rec_table.columns[0].width = Inches(5.0)
        rec_table.columns[1].width = Inches(2.0)
        
        for item in recurring_costs:
            row_cells = rec_table.add_row().cells
            for cell in row_cells:
                set_cell_background(cell, "FFFFFF")
                set_cell_margins(cell, 80, 80, 100, 100)
                set_cell_bottom_border(cell, "E2E8F0", "4")
            
            p_desc = row_cells[0].paragraphs[0]
            run_lbl = p_desc.add_run(f"Estimation mensuelle post-lancement : {item.get('designation', '')}")
            run_lbl.font.size = Pt(8.5)
            run_lbl.font.color.rgb = c_light
            
            p_val = row_cells[1].paragraphs[0]
            p_val.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run_val = p_val.add_run(f"{item.get('price', '')} / {item.get('period', 'mois')}")
            run_val.bold = True
            run_val.font.size = Pt(8.5)
            
        doc.add_paragraph().paragraph_format.space_after = Pt(15)
        
    summary_table = doc.add_table(rows=0, cols=2)
    summary_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    summary_table.autofit = False
    summary_table.columns[0].width = Inches(2.2)
    summary_table.columns[1].width = Inches(1.5)
    
    doc_type_label = "Développement" if doc_type == "DEVIS" else "Facture"
    row = summary_table.add_row()
    row.cells[0].text = f"Total {doc_type_label} HT"
    row.cells[1].text = format_currency(total_ht, currency)
    set_cell_bottom_border(row.cells[0], "E2E8F0", "4")
    set_cell_bottom_border(row.cells[1], "E2E8F0", "4")
    
    tax_rate = float(yaml_data.get("tax_rate", 0.0))
    if tax_rate > 0:
        row = summary_table.add_row()
        row.cells[0].text = f"TVA ({int(tax_rate*100)}%)"
        row.cells[1].text = format_currency(vat_amount, currency)
        set_cell_bottom_border(row.cells[0], "E2E8F0", "4")
        set_cell_bottom_border(row.cells[1], "E2E8F0", "4")
        
    for r in summary_table.rows:
        for cell in r.cells:
            set_cell_margins(cell, 60, 60, 80, 80)
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(9.5)
            p.runs[0].font.bold = True
        r.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
    div_row = summary_table.add_row()
    set_cell_background(div_row.cells[0], "103824")
    set_cell_background(div_row.cells[1], "103824")
    set_cell_margins(div_row.cells[0], 20, 20, 0, 0)
    set_cell_margins(div_row.cells[1], 20, 20, 0, 0)
    
    ttc_row = summary_table.add_row()
    ttc_row.cells[0].text = "Total Général TTC"
    total_val_s = f"{int(round(total_ttc)):,}".replace(",", " ")
    ttc_row.cells[1].text = f"{total_val_s}\n{currency}"
    
    for cell in ttc_row.cells:
        set_cell_background(cell, "FFFFFF")
        set_cell_margins(cell, 80, 80, 80, 80)
        p = cell.paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(12)
        p.runs[0].font.color.rgb = c_forest
        
    ttc_row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if len(ttc_row.cells[1].paragraphs[0].runs) > 0:
         run_cur = ttc_row.cells[1].paragraphs[0].runs[0]
         run_cur.font.size = Pt(12)
         
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    
    terms = yaml_data.get("terms", [])
    if terms:
        terms_table = doc.add_table(rows=1, cols=1)
        terms_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        terms_table.columns[0].width = Inches(7.0)
        t_cell = terms_table.cell(0, 0)
        set_cell_background(t_cell, "FAFAFA")
        set_cell_borders(t_cell, "E2E8F0", "E2E8F0", "E2E8F0", "E2E8F0", "4")
        set_cell_margins(t_cell, 140, 140, 180, 180)
        
        p_t_hdr = t_cell.paragraphs[0]
        p_t_hdr.paragraph_format.space_after = Pt(6)
        run_t_hdr = p_t_hdr.add_run("Conditions de paiement & Modalités :")
        run_t_hdr.bold = True
        run_t_hdr.font.size = Pt(9)
        run_t_hdr.font.color.rgb = c_dark
        
        for term in terms:
            p_term = t_cell.add_paragraph()
            p_term.paragraph_format.left_indent = Inches(0.15)
            p_term.paragraph_format.space_after = Pt(3)
            run_term = p_term.add_run(f"• {term}")
            run_term.font.size = Pt(8.5)
            run_term.font.color.rgb = c_dark
            
    doc.add_paragraph().paragraph_format.space_after = Pt(25)
    
    # Signature Boxes with solid top borders matching the layout
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False
    sig_table.columns[0].width = Inches(3.3)
    sig_table.columns[1].width = Inches(3.7)
    
    c_sig_client = sig_table.cell(0, 0)
    c_sig_agency = sig_table.cell(0, 1)
    
    # Draw solid top border for both signature boxes
    set_cell_borders(c_sig_client, top="E2E8F0", bottom=None, left=None, right=None, size="4")
    set_cell_borders(c_sig_agency, top="E2E8F0", bottom=None, left=None, right=None, size="4")
    
    set_cell_margins(c_sig_client, 100, 100, 100, 100)
    set_cell_margins(c_sig_agency, 100, 100, 100, 100)
    
    p_sig_c = c_sig_client.paragraphs[0]
    p_sig_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sig_c.paragraph_format.space_before = Pt(30)
    p_sig_c.paragraph_format.space_after = Pt(2)
    run_sig_c_hdr = p_sig_c.add_run("Bon pour accord\n")
    run_sig_c_hdr.font.size = Pt(8.5)
    run_sig_c_lbl = p_sig_c.add_run("Signature du Client (précédée de la date)")
    run_sig_c_lbl.font.size = Pt(8.5)
    
    p_sig_a = c_sig_agency.paragraphs[0]
    p_sig_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sig_a.paragraph_format.space_before = Pt(30)
    p_sig_a.paragraph_format.space_after = Pt(2)
    run_sig_a_lbl = p_sig_a.add_run("Cachet et Signature de l'Agence")
    run_sig_a_lbl.font.size = Pt(8.5)
    
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.space_before = Pt(30)
    run_foot = p_foot.add_run(f"{company_info.get('name', 'Iboga Lab')} SARL - Capital Social de {company_info.get('capital', '1 000 000 FCFA')} - {company_info.get('address', 'Libreville, Gabon')} - {company_info.get('website', 'www.ibogalab.tech')}")
    run_foot.font.size = Pt(7.5)
    run_foot.font.color.rgb = c_light
    
    doc.save(docx_path)

def watch_file(md_path):
    print(f"Watching {md_path} for changes... (Press Ctrl+C to stop)")
    last_mtime = os.path.getmtime(md_path)
    
    compile_document(md_path)
    
    try:
        while True:
            time.sleep(1)
            if not os.path.exists(md_path):
                continue
            current_mtime = os.path.getmtime(md_path)
            if current_mtime > last_mtime:
                print("Change detected, regenerating...")
                time.sleep(0.5)
                compile_document(md_path)
                last_mtime = current_mtime
    except KeyboardInterrupt:
        print("\nStopped watching.")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python sync.py [--watch] <path_to_markdown>")
        sys.exit(1)
        
    watch_mode = False
    file_path = ""
    
    for arg in sys.argv[1:]:
        if arg == "--watch":
            watch_mode = True
        else:
            file_path = arg
            
    if not file_path:
        print("Error: Missing markdown file path.")
        sys.exit(1)
        
    if not os.path.exists(file_path):
        print(f"Error: File not found: {file_path}")
        sys.exit(1)
        
    if watch_mode:
        watch_file(file_path)
    else:
        compile_document(file_path)
